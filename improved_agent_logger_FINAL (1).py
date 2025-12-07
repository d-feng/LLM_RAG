from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import re
import json
from typing import Dict, Any


def clean_ansi_codes(text: str) -> str:
    """Remove ANSI escape codes from text"""
    ansi_escape = re.compile(r'\x1b\[[0-9;]*m')
    return ansi_escape.sub('', text)


def format_complex_data(text: str) -> str:
    """
    Format complex data structures (dicts, JSON) to be human-readable.
    Handles Python dict strings like those from cBioPortal queries.
    """
    import ast
    
    # Check if this looks like it contains a Python dict
    if not ('{' in text and ':' in text):
        return text
    
    # Try to find and parse dict-like patterns
    result_parts = []
    remaining_text = text
    
    # Look for patterns like "label: {dict}" or standalone dicts
    while remaining_text:
        # Try to find a dict in the text
        dict_start = remaining_text.find('{')
        if dict_start == -1:
            # No more dicts, add remaining text
            if remaining_text.strip():
                result_parts.append(remaining_text.strip())
            break
        
        # Add any text before the dict
        before_dict = remaining_text[:dict_start].strip()
        if before_dict and not before_dict.endswith(':'):
            result_parts.append(before_dict)
        elif before_dict:
            # This is a label like "cBioPortal query result:"
            result_parts.append(before_dict)
        
        # Try to extract and parse the dict
        try:
            # Find the matching closing brace
            brace_count = 0
            dict_end = dict_start
            in_string = False
            escape_next = False
            
            for i in range(dict_start, len(remaining_text)):
                char = remaining_text[i]
                
                if escape_next:
                    escape_next = False
                    continue
                
                if char == '\\':
                    escape_next = True
                    continue
                
                if char in ('"', "'"):
                    in_string = not in_string
                
                if not in_string:
                    if char == '{':
                        brace_count += 1
                    elif char == '}':
                        brace_count -= 1
                        if brace_count == 0:
                            dict_end = i + 1
                            break
            
            # Extract the dict string
            dict_str = remaining_text[dict_start:dict_end]
            
            # Try to parse it
            try:
                data = ast.literal_eval(dict_str)
                # Successfully parsed - format it
                formatted = format_dict_human_readable(data)
                result_parts.append(formatted)
                
                # Continue with remaining text
                remaining_text = remaining_text[dict_end:].strip()
            except (ValueError, SyntaxError):
                # Couldn't parse - just include the raw text
                result_parts.append(dict_str[:200] + "..." if len(dict_str) > 200 else dict_str)
                remaining_text = remaining_text[dict_end:].strip()
        
        except Exception:
            # Something went wrong, just add a portion and continue
            result_parts.append(remaining_text[:100] + "...")
            break
    
    return '\n\n'.join(result_parts)


def format_dict_human_readable(data, indent=0, max_depth=3) -> str:
    """
    Convert a dictionary/list to human-readable format
    Optimized for database query results like cBioPortal
    """
    if indent > max_depth:
        return "{...}" if isinstance(data, dict) else "[...]"
    
    if data is None:
        return "None"
    
    if isinstance(data, bool):
        return "✓ Yes" if data else "✗ No"
    
    if isinstance(data, (int, float)):
        return str(data)
    
    if isinstance(data, str):
        # Limit very long strings
        if len(data) > 150:
            return data[:150] + "..."
        return data
    
    if isinstance(data, list):
        if not data:
            return "(empty list)"
        
        # For lists of dicts (like query results), show summary
        if all(isinstance(x, dict) for x in data):
            result = [f"📋 Found {len(data)} result(s):"]
            for i, item in enumerate(data[:5], 1):  # Show first 5 items
                result.append(f"\n  Result {i}:")
                # Show key fields from each result
                formatted_item = format_dict_human_readable(item, indent + 1, max_depth)
                for line in formatted_item.split('\n'):
                    result.append(f"    {line}")
            
            if len(data) > 5:
                result.append(f"\n  ... and {len(data) - 5} more results")
            
            return ''.join(result)
        
        # For simple lists
        if len(data) <= 5 and all(isinstance(x, (str, int, float, bool)) for x in data):
            return ", ".join(str(x) for x in data)
        
        # For other lists
        result = []
        for i, item in enumerate(data[:5], 1):
            formatted = format_dict_human_readable(item, indent + 1, max_depth)
            if '\n' in formatted:
                result.append(f"\n  {i}. {formatted}")
            else:
                result.append(f"\n  {i}. {formatted}")
        
        if len(data) > 5:
            result.append(f"\n  ... and {len(data) - 5} more items")
        
        return ''.join(result)
    
    if isinstance(data, dict):
        if not data:
            return "(empty dict)"
        
        result = []
        
        # Key fields to show prominently
        priority_keys = [
            'success', 'name', 'description', 'endpoint', 'method',
            'studyId', 'cancerTypeId', 'molecularAlterationType',
            'datatype', 'molecularProfileId', 'citation', 'pmid'
        ]
        
        # Secondary keys that are useful but less critical
        secondary_keys = [
            'showProfileInAnalysisTab', 'patientLevel', 'publicStudy',
            'referenceGenome', 'importDate', 'status'
        ]
        
        # Keys to skip (too technical or redundant)
        skip_keys = ['readPermission', 'groups']
        
        # Organize keys
        shown_keys = set()
        
        # Show priority keys first
        for key in priority_keys:
            if key in data:
                value = data[key]
                formatted_value = format_dict_human_readable(value, indent + 1, max_depth)
                
                # Format based on key type
                if key == 'success':
                    result.append(f"✓ Status: {formatted_value}")
                elif key == 'endpoint':
                    # Shorten long URLs
                    if len(str(value)) > 80:
                        result.append(f"🔗 API Endpoint: {str(value)[:80]}...")
                    else:
                        result.append(f"🔗 API Endpoint: {value}")
                elif key in ['studyId', 'molecularProfileId']:
                    result.append(f"🔬 {key}: {value}")
                elif key == 'name':
                    result.append(f"📌 Name: {formatted_value}")
                elif key == 'description':
                    result.append(f"📝 Description: {formatted_value}")
                elif key == 'method':
                    result.append(f"⚙️ Method: {value}")
                elif key == 'cancerTypeId':
                    result.append(f"🧬 Cancer Type: {value}")
                elif key == 'citation':
                    result.append(f"📚 Citation: {value}")
                elif key == 'pmid':
                    result.append(f"📄 PMID: {value}")
                else:
                    result.append(f"{key}: {formatted_value}")
                
                shown_keys.add(key)
        
        # Show nested structures (like 'result', 'query_info', 'study')
        for key in data:
            if key in shown_keys or key in skip_keys:
                continue
            
            value = data[key]
            
            # Handle nested dicts/lists specially
            if isinstance(value, (dict, list)) and key not in secondary_keys:
                formatted_value = format_dict_human_readable(value, indent + 1, max_depth)
                if key == 'result':
                    result.append(f"\n📊 Query Results:")
                    result.append(formatted_value)
                elif key == 'query_info':
                    result.append(f"\n🔍 Query Information:")
                    for line in formatted_value.split('\n'):
                        if line.strip():
                            result.append(f"  {line}")
                elif key == 'study':
                    if indent < 2:  # Only show study details at top levels
                        result.append(f"\n  Study Info: {formatted_value}")
                else:
                    result.append(f"{key}: {formatted_value}")
                shown_keys.add(key)
        
        # Show secondary keys if at top level
        if indent == 0:
            secondary_present = [k for k in secondary_keys if k in data and k not in shown_keys]
            if secondary_present:
                sec_values = []
                for key in secondary_present:
                    val = format_dict_human_readable(data[key], indent + 1, max_depth)
                    sec_values.append(f"{key}={val}")
                    shown_keys.add(key)
                if sec_values:
                    result.append(f"\n  Additional: {', '.join(sec_values)}")
        
        # Count remaining keys
        remaining = len(data) - len(shown_keys)
        if remaining > 0:
            result.append(f"\n  ... and {remaining} more field(s)")
        
        return '\n'.join(result)
    
    return str(data)


def normalize_line_breaks(text: str) -> str:
    """
    Convert both \\n (escaped) and \n (actual) to consistent newlines
    Handles cases where text has been double-escaped or more
    """
    # Handle multiple levels of escaping - keep replacing until no more \\n
    while '\\n' in text:
        text = text.replace('\\n', '\n')
    
    # Also handle other escaped characters
    while '\\t' in text:
        text = text.replace('\\t', '\t')
    
    while '\\r' in text:
        text = text.replace('\\r', '\r')
    
    return text


def add_text_with_line_breaks(paragraph, text: str):
    """
    Add text to paragraph with proper line breaks instead of showing \\n
    """
    lines = text.split('\n')
    for i, line in enumerate(lines):
        paragraph.add_run(line)
        if i < len(lines) - 1:  # Don't add break after last line
            paragraph.add_run('\n')


def add_paragraph_with_line_breaks(doc, text: str, style=None):
    """
    Add a paragraph with text that has proper line breaks
    """
    para = doc.add_paragraph(style=style)
    add_text_with_line_breaks(para, text)
    return para


def extract_agent_messages(output) -> list:
    """
    Extract and parse agent messages from output
    Returns list of dicts with message type and content
    """
    messages = []
    
    # Convert output to string
    output_str = str(output)
    
    # Clean ANSI codes
    output_str = clean_ansi_codes(output_str)
    
    # Normalize line breaks (convert \\n to \n)
    output_str = normalize_line_breaks(output_str)
    
    # Remove list/tuple wrapper if present
    if output_str.startswith("(['") or output_str.startswith('(["'):
        # Extract content between quotes
        output_str = output_str[2:-2] if output_str.endswith("'])") or output_str.endswith('"])') else output_str[2:]
    
    # Split by message separators
    message_pattern = r'={32,}\s*(Human Message|Ai Message)\s*={32,}'
    parts = re.split(message_pattern, output_str)
    
    current_type = None
    for i, part in enumerate(parts):
        part = part.strip()
        
        if 'Human Message' in part:
            current_type = 'human'
        elif 'Ai Message' in part:
            current_type = 'ai'
        elif part and current_type:
            messages.append({
                'type': current_type,
                'content': part
            })
            current_type = None
    
    return messages


def extract_code_and_observations(text: str) -> list:
    """Extract code blocks and observations from text"""
    blocks = []
    
    # Pattern for execute blocks
    execute_pattern = r'<execute>(.*?)</execute>'
    observation_pattern = r'<observation>(.*?)</observation>'
    
    # Find all execute blocks
    for match in re.finditer(execute_pattern, text, re.DOTALL):
        blocks.append({
            'type': 'code',
            'content': match.group(1).strip()
        })
    
    # Find all observation blocks
    for match in re.finditer(observation_pattern, text, re.DOTALL):
        blocks.append({
            'type': 'observation',
            'content': match.group(1).strip()
        })
    
    return blocks


def extract_solution(text: str) -> str:
    """Extract content from solution tags"""
    match = re.search(r'<solution>(.*?)</solution>', text, re.DOTALL)
    if match:
        return match.group(1).strip()
    return None


def add_colored_box(doc, text: str, color_rgb: tuple, heading: str = None):
    """Add a colored text box to document with proper line breaks"""
    if heading:
        para = doc.add_paragraph()
        run = para.add_run(heading)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(*color_rgb)
    
    # Create paragraph with line breaks
    para = doc.add_paragraph()
    add_text_with_line_breaks(para, text)
    
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.right_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    
    # Add shading
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), f"{color_rgb[0]:02x}{color_rgb[1]:02x}{color_rgb[2]:02x}20")
    para._element.get_or_add_pPr().append(shading_elm)
    
    return para


def log_agent_results_readable(results: Dict[str, Any], 
                               filename: str = None,
                               title: str = "Agent Execution Report",
                               include_code: bool = True,
                               include_observations: bool = True):
    """
    Create a human-readable Word document from agent chain results
    
    Parameters:
    -----------
    results : dict
        Results from connector.execute_chain()
    filename : str
        Output filename (auto-generated if None)
    title : str
        Document title
    include_code : bool
        Whether to include code blocks
    include_observations : bool
        Whether to include observation blocks
    """
    
    # Generate filename
    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"agent_report_{timestamp}.docx"
    
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========== TITLE ==========
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========== EXECUTIVE SUMMARY ==========
    doc.add_heading('Executive Summary', level=1)
    
    total_time = sum(r['execution_time'] for r in results.values())
    total_steps = len(results)
    
    # Create summary table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    summary_data = [
        ('Execution Date', datetime.now().strftime("%B %d, %Y at %I:%M %p")),
        ('Total Steps', str(total_steps)),
        ('Total Time', f"{total_time:.1f} seconds ({total_time/60:.1f} minutes)"),
        ('Average per Step', f"{total_time/total_steps:.1f} seconds")
    ]
    
    for i, (label, value) in enumerate(summary_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # ========== AGENT FLOW ==========
    doc.add_heading('Agent Processing Flow', level=1)
    
    sorted_results = sorted(results.items(), key=lambda x: x[1]['step'])
    flow_items = []
    for agent_name, result in sorted_results:
        flow_items.append(f"{result['step']}. {agent_name.replace('_', ' ').title()}")
    
    # Create flow paragraph with proper line breaks
    flow_text = '\n'.join(flow_items)
    add_paragraph_with_line_breaks(doc, flow_text)
    
    doc.add_page_break()
    
    # ========== DETAILED RESULTS ==========
    doc.add_heading('Detailed Agent Results', level=1)
    
    for agent_name, result in sorted_results:
        agent = result['agent']
        
        # Agent header
        doc.add_heading(f"Step {result['step']}: {agent_name.replace('_', ' ').title()}", level=2)
        
        # Agent info box with proper line breaks
        info_text = f"Role: {agent.description}\nExecution Time: {result['execution_time']:.1f} seconds"
        info_para = doc.add_paragraph()
        
        # Split by newline and format each part
        parts = info_text.split('\n')
        for i, part in enumerate(parts):
            if part.startswith('Role: '):
                info_para.add_run('Role: ').bold = True
                info_para.add_run(part[6:])  # Text after "Role: "
            elif part.startswith('Execution Time: '):
                if i > 0:  # Add line break before if not first
                    info_para.add_run('\n')
                info_para.add_run('Execution Time: ').bold = True
                info_para.add_run(part[16:])  # Text after "Execution Time: "
            else:
                if i > 0:
                    info_para.add_run('\n')
                info_para.add_run(part)
        
        doc.add_paragraph()
        
        # ========== INPUT SECTION ==========
        doc.add_heading('Input Query', level=3)
        
        input_text = str(result['input'])
        input_text = clean_ansi_codes(input_text)
        input_text = normalize_line_breaks(input_text)  # Fix \\n to \n
        
        # Check if input is from previous agent's solution
        solution_in_input = extract_solution(input_text)
        if solution_in_input:
            add_colored_box(doc, solution_in_input, (76, 175, 80))  # Green
        else:
            # Add paragraph with proper line breaks
            add_paragraph_with_line_breaks(doc, input_text)
        
        doc.add_paragraph()
        
        # ========== OUTPUT SECTION ==========
        doc.add_heading('Agent Response', level=3)
        
        # Parse messages from output
        messages = extract_agent_messages(result['output'])
        
        if not messages:
            # Fallback: just clean and display the output
            output_text = clean_ansi_codes(str(result['output']))
            output_text = normalize_line_breaks(output_text)  # Fix \\n to \n
            add_paragraph_with_line_breaks(doc, output_text)
        else:
            for msg in messages:
                if msg['type'] == 'human':
                    add_colored_box(doc, msg['content'], (33, 150, 243), "Human Message:")
                elif msg['type'] == 'ai':
                    # Extract solution if present
                    solution = extract_solution(msg['content'])
                    
                    if solution:
                        # Show solution prominently
                        doc.add_heading('Solution:', level=4)
                        solution_para = doc.add_paragraph(style='Intense Quote')
                        add_text_with_line_breaks(solution_para, solution)
                    else:
                        # Show regular AI message
                        # Check for code and observations
                        blocks = extract_code_and_observations(msg['content'])
                        
                        if blocks:
                            # Display structured content
                            remaining_text = msg['content']
                            
                            # Remove code and observation blocks from remaining text
                            for block in blocks:
                                if block['type'] == 'code':
                                    remaining_text = re.sub(
                                        r'<execute>.*?</execute>', '', 
                                        remaining_text, 
                                        flags=re.DOTALL
                                    )
                                else:
                                    remaining_text = re.sub(
                                        r'<observation>.*?</observation>', '', 
                                        remaining_text, 
                                        flags=re.DOTALL
                                    )
                            
                            # Display remaining text first
                            remaining_text = remaining_text.strip()
                            if remaining_text and len(remaining_text) > 10:
                                add_paragraph_with_line_breaks(doc, remaining_text)
                            
                            # Display code blocks
                            if include_code:
                                for block in blocks:
                                    if block['type'] == 'code':
                                        doc.add_heading('Code Executed:', level=4)
                                        code_para = doc.add_paragraph(style='No Spacing')
                                        add_text_with_line_breaks(code_para, block['content'])
                                        for run in code_para.runs:
                                            run.font.name = 'Courier New'
                                            run.font.size = Pt(9)
                            
                            # Display observations
                            if include_observations:
                                for block in blocks:
                                    if block['type'] == 'observation':
                                        doc.add_heading('Observation:', level=4)
                                        obs_text = block['content']
                                        
                                        # Format complex data structures
                                        obs_text = format_complex_data(obs_text)
                                        
                                        # Limit length
                                        if len(obs_text) > 2000:
                                            obs_text = obs_text[:2000] + "\n\n... (output truncated for readability)"
                                        
                                        add_colored_box(doc, obs_text, (255, 193, 7))  # Amber
                        else:
                            # Just display the content
                            content = msg['content'].strip()
                            if content:
                                add_paragraph_with_line_breaks(doc, content)
        
        # Add separator between agents
        doc.add_paragraph('─' * 80)
        doc.add_page_break()
    
    # ========== FINAL SOLUTION ==========
    # Check if last agent has a solution
    if sorted_results:
        final_output = str(sorted_results[-1][1]['output'])
        final_output = normalize_line_breaks(final_output)  # Fix \\n to \n
        final_solution = extract_solution(final_output)
        
        if final_solution:
            doc.add_heading('Final Solution', level=1)
            
            # Add solution in a nice format with line breaks
            solution_para = doc.add_paragraph()
            add_text_with_line_breaks(solution_para, final_solution)
            solution_para.paragraph_format.left_indent = Inches(0.5)
            solution_para.paragraph_format.right_indent = Inches(0.5)
            
            # Add green shading
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D4EDDA')  # Light green
            solution_para._element.get_or_add_pPr().append(shading_elm)
    
    # Save document - use proper path
    if not filename.startswith('/'):
        output_path = f"./{filename}"
    else:
        output_path = filename
    
    doc.save(output_path)
    
    print(f"✅ Readable agent report created: {filename}")
    print(f"   📊 {total_steps} steps completed")
    print(f"   ⏱️  Total time: {total_time:.1f}s ({total_time/60:.1f} min)")
    print(f"   📄 Location: {output_path}")
    
    return output_path


# Simplified version with minimal formatting
def log_agent_results_simple(results: Dict[str, Any], filename: str = None):
    """
    Create a simple, clean Word document - just the essentials
    """
    
    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"agent_simple_{timestamp}.docx"
    
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    doc = Document()
    
    # Title
    doc.add_heading('Agent Chain Results', 0)
    
    # Summary
    total_time = sum(r['execution_time'] for r in results.values())
    doc.add_paragraph(f"Completed {len(results)} steps in {total_time:.1f} seconds")
    doc.add_paragraph()
    
    # Process each agent
    sorted_results = sorted(results.items(), key=lambda x: x[1]['step'])
    
    for agent_name, result in sorted_results:
        # Step header
        doc.add_heading(f"Step {result['step']}: {agent_name.upper()}", level=2)
        
        # Input
        doc.add_heading('Input:', level=3)
        input_text = clean_ansi_codes(str(result['input']))
        input_text = normalize_line_breaks(input_text)  # Fix \\n to \n
        solution = extract_solution(input_text)
        add_paragraph_with_line_breaks(doc, solution if solution else input_text)
        
        # Output - extract just the solution or final message
        doc.add_heading('Output:', level=3)
        output_text = clean_ansi_codes(str(result['output']))
        output_text = normalize_line_breaks(output_text)  # Fix \\n to \n
        solution = extract_solution(output_text)
        
        if solution:
            add_paragraph_with_line_breaks(doc, solution)
        else:
            # Get last AI message
            messages = extract_agent_messages(result['output'])
            if messages:
                last_ai = [m for m in messages if m['type'] == 'ai']
                if last_ai:
                    # Clean up the content
                    content = last_ai[-1]['content']
                    # Remove code/observation blocks for simplicity
                    content = re.sub(r'<execute>.*?</execute>', '', content, flags=re.DOTALL)
                    content = re.sub(r'<observation>.*?</observation>', '', content, flags=re.DOTALL)
                    content = content.strip()
                    if content:
                        add_paragraph_with_line_breaks(doc, content)
        
        doc.add_paragraph()
    
    # Save - use proper path
    if not filename.startswith('/'):
        output_path = f"./{filename}"
    else:
        output_path = filename
    
    doc.save(output_path)
    
    print(f"✅ Simple report created: {filename}")
    print(f"   📄 Location: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Agent logging functions loaded!")
    print("\nAvailable functions:")
    print("  - log_agent_results_readable() : Full formatted report")
    print("  - log_agent_results_simple()   : Clean, minimal report")
    print("\nKey features:")
    print("  ✅ Converts \\n to actual line breaks in Word")
    print("  ✅ Removes ANSI escape codes")
    print("  ✅ Extracts and highlights solutions")
    print("  ✅ Professional formatting")
